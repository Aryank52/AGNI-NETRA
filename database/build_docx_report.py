"""
Generates the publication-quality DOCX version of the AGNI-NETRA Comprehensive Research Project Report.
Uses python-docx to build a professionally styled Word document with:
- Title Page / Cover
- Numbered Headings & Hierarchy
- Formatted Tables with Header Shading & Borders
- Callout Alert Boxes
- Header & Footer with Page Numbering
- Mathematical Formulations
- Complete 31 Sections + Abstract, Keywords, References, Appendices
"""

import sys
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

ROOT_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DOCX = ROOT_DIR / "AGNI_NETRA_COMPLETE_RESEARCH_PROJECT_REPORT.docx"


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def add_callout(doc, text, title="NOTE", fill_hex="F0F4F8", border_hex="003366"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border only
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"[{title}] ")
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(0, 51, 102)
    r_text = p.add_run(text)
    r_text.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def format_table(tbl, col_widths, headers, data, header_bg="003366", alt_bg="F8F9FA"):
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header Row
    hdr_cells = tbl.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], header_bg)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = "Calibri"

    # Data Rows
    for r_idx, row in enumerate(data):
        row_cells = tbl.add_row().cells
        bg = alt_bg if (r_idx % 2 == 1) else "FFFFFF"
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=140, right=140)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = "Calibri"

    # Widths
    for row in tbl.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = Inches(width)


def build_docx():
    print(f"Generating DOCX report: {OUTPUT_DOCX}...")
    doc = Document()

    # Configure Margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
        # Add Header & Footer
        header = s.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("AGNI-NETRA: Geospatial Intelligence & ML Platform for Thermal Risk Assessment")
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(128, 128, 128)

        footer = s.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("AGNI-NETRA Research Project Report — Production Release v1.0 — Controlled Activation")
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(128, 128, 128)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(30, 30, 30)

    # COVER PAGE
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(12)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("AGNI-NETRA:\nA Geospatial Intelligence and Machine-Learning Platform for Satellite-Based Thermal Event Detection, Contextualization, Risk Assessment, and Human-in-the-Loop Decision Support in India")
    r_title.bold = True
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(0, 51, 102)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(36)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("From Raw Satellite Thermal Anomaly Ingestion to Multi-Source Contextualization, Calibrated Risk Scoring, and National Command Center Operations")
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(90, 90, 90)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(48)
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("AGNI-NETRA Project Core Engineering Team\nAdvanced Geospatial & Thermal Intelligence Initiative, India\nSeptember 2026 | Production Go-Live Release (Phase 16)")
    r_meta.font.size = Pt(11)
    r_meta.font.color.rgb = RGBColor(60, 60, 60)

    # Executive Status Table on Cover
    cov_tbl = doc.add_table(rows=1, cols=2)
    cov_data = [
        ["Database Engine", "PostgreSQL 16.1 + PostGIS 3.4.2 (EPSG:4326)"],
        ["Authoritative Thermal Records", "8,221,825+ Observations (2022–2026 Archive)"],
        ["Historical Partition Status", "100% Sealed & Immutable (6,448,666 rows, 0 diff)"],
        ["Geospatial Knowledge Base", "35,684 Facilities | 1,633 CEA Plants | 98,793 Mines"],
        ["Champion Model", "xgb-v3.0-real-candidate + Balanced Platt Calibrator"],
        ["Spatial Validation Accuracy", "94.32% Mean Accuracy (Macro F1: 0.9318)"],
        ["Probability Calibration", "Log-Loss: 0.7124 (55.7% drop) | ECE: 0.1294 (54.3% drop)"],
        ["Tri-Tier Decision Safety", "Tier 1 Selective Accuracy: 97.18%"],
        ["Operational Dispatch Gate", "ENABLE_OPERATIONAL_DISPATCH_GATE = False (Secured)"],
        ["Platform Operational Status", "HEALTHY / PRODUCTION-READY (CONTROLLED ACTIVATION)"]
    ]
    format_table(cov_tbl, [2.8, 4.0], ["Core Platform Dimension", "Verified Operational State"], cov_data, header_bg="003366")

    doc.add_page_break()

    # Read Markdown Content and Convert Systematically
    md_text = (ROOT_DIR / "AGNI_NETRA_COMPLETE_RESEARCH_PROJECT_REPORT.md").read_text(encoding="utf-8")
    lines = md_text.splitlines()

    in_abstract = False
    in_table = False
    table_headers = []
    table_rows = []
    col_widths = []

    for line in lines:
        sline = line.strip()

        # Skip main markdown title as we have a cover page
        if sline.startswith("# AGNI-NETRA:"):
            continue
        if sline.startswith("**Subtitle**:") or sline.startswith("**Authors") or sline.startswith("**Date") or sline.startswith("**Document") or sline.startswith("**Repository") or sline.startswith("**Software"):
            continue
        if sline == "---":
            continue

        # Headings
        if sline.startswith("## Abstract"):
            in_abstract = True
            h = doc.add_heading("Abstract", level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
            continue
        elif sline.startswith("## Keywords"):
            in_abstract = False
            h = doc.add_heading("Keywords", level=1)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            continue
        elif sline.startswith("## ") and not sline.startswith("### "):
            title = sline.replace("## ", "").strip()
            h = doc.add_heading(title, level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
            continue
        elif sline.startswith("### "):
            title = sline.replace("### ", "").strip()
            h = doc.add_heading(title, level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            continue
        elif sline.startswith("#### "):
            title = sline.replace("#### ", "").strip()
            h = doc.add_heading(title, level=3)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            continue

        # Callouts / Notes
        if sline.startswith("> **Key Result**:") or sline.startswith("> **Scientific Finding**:") or sline.startswith("> **Safety Significance**:") or sline.startswith("> [!NOTE]"):
            callout_text = sline.replace("> **Key Result**:", "").replace("> **Scientific Finding**:", "").replace("> **Safety Significance**:", "").replace("> [!NOTE]", "").replace(">", "").strip()
            add_callout(doc, callout_text, title="KEY FINDING", fill_hex="EBF3FA", border_hex="003366")
            continue

        # Tables Parsing
        if sline.startswith("|") and sline.endswith("|"):
            cells = [c.strip() for c in sline.split("|")[1:-1]]
            if not cells:
                continue
            if all(set(c).issubset({'-', ':', ' '}) for c in cells):
                # Separator line
                in_table = True
                continue
            if not in_table:
                # Header row
                table_headers = [c.replace("**", "").replace("`", "") for c in cells]
                table_rows = []
            else:
                # Data row
                table_rows.append([c.replace("**", "").replace("`", "") for c in cells])
            continue
        elif in_table and not sline.startswith("|"):
            # End of table, render it!
            in_table = False
            if table_headers and table_rows:
                num_cols = len(table_headers)
                t = doc.add_table(rows=1, cols=num_cols)
                # Compute balanced widths
                total_w = 6.8
                w_per_col = total_w / num_cols
                widths = [w_per_col] * num_cols
                if num_cols == 2:
                    widths = [2.5, 4.3]
                elif num_cols == 3:
                    widths = [2.2, 2.3, 2.3]
                elif num_cols == 4:
                    widths = [1.8, 1.6, 1.6, 1.8]
                elif num_cols == 5:
                    widths = [1.5, 1.3, 1.3, 1.3, 1.4]
                elif num_cols == 6:
                    widths = [1.3, 1.1, 1.1, 1.1, 1.1, 1.1]

                format_table(t, widths, table_headers, table_rows)
                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_after = Pt(8)
                table_headers = []
                table_rows = []

        # Code blocks & equations
        if sline.startswith("```") or sline.endswith("```"):
            continue
        if sline.startswith("$$") and sline.endswith("$$"):
            p_eq = doc.add_paragraph()
            p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_eq = p_eq.add_run(sline.replace("$$", "").strip())
            r_eq.font.name = "Cambria Math"
            r_eq.font.size = Pt(11)
            r_eq.font.italic = True
            p_eq.paragraph_format.space_before = Pt(4)
            p_eq.paragraph_format.space_after = Pt(6)
            continue

        # Regular Paragraph
        if sline:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            # Bold & inline formatting
            clean_line = sline.replace("**", "").replace("`", "").replace("$", "")
            if in_abstract:
                r = p.add_run(clean_line)
                r.font.size = Pt(10.5)
                r.font.italic = True
                p.paragraph_format.line_spacing = 1.2
            else:
                r = p.add_run(clean_line)
                r.font.size = Pt(10.5)

    doc.save(str(OUTPUT_DOCX))
    print(f"Successfully generated Word report: {OUTPUT_DOCX}")


if __name__ == "__main__":
    build_docx()
